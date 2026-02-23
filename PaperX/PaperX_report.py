from email.mime import text
import os
import re
import unicodedata
import shutil
from dataclasses import dataclass
from docx import Document

import tempfile
from copy import deepcopy
from shutil import which

import subprocess
from pathlib import Path



# ======================================================================
# Word OMML (Equation Editor) support: OMML -> LaTeX (pandoc) -> image fallback
# ======================================================================

def has_omml(paragraph) -> bool:
    """
    Word 'Insert Equation' creates OMML nodes (m:oMath / m:oMathPara).
    python-docx paragraph.text does NOT include them.
    """
    xml = paragraph._element.xml
    return ("<m:oMath" in xml) or ("<m:oMathPara" in xml)

def write_single_paragraph_docx(src_paragraph, out_docx_path: str):
    """
    src_paragraph._element (including OMML) is deep-copied into a new docx.
    This preserves the equation XML so pandoc can see it.
    """
    d = Document()
    # remove default empty paragraph
    if d.paragraphs:
        p0 = d.paragraphs[0]._element
        p0.getparent().remove(p0)
    d._body._element.append(deepcopy(src_paragraph._element))
    d.save(out_docx_path)

def extract_math_from_pandoc_latex(latex_fragment: str) -> str | None:
    """
    pandoc output may look like:
      \\[ ... \\]
      $$ ... $$
      \\( ... \\)
    We extract inner math for inserting into \\begin{equation} ... \\end{equation}
    """
    if not latex_fragment:
        return None

    s = latex_fragment.strip()

    # common wrappers
    # \[ ... \]
    m = re.search(r"\\\[(.*)\\\]", s, flags=re.S)
    if m:
        inner = m.group(1).strip()
        return inner if inner else None

    # $$ ... $$
    m = re.search(r"\$\$(.*)\$\$", s, flags=re.S)
    if m:
        inner = m.group(1).strip()
        return inner if inner else None

    # \( ... \)
    m = re.search(r"\\\((.*)\\\)", s, flags=re.S)
    if m:
        inner = m.group(1).strip()
        return inner if inner else None

    # if no wrapper found, return as-is (sometimes pandoc outputs bare math)
    return s if s else None

def pandoc_docx_paragraph_to_latex_math(src_paragraph) -> str | None:
    """
    Converts a single paragraph docx (with OMML) to LaTeX using pandoc,
    and extracts the math content.
    """
    if which("pandoc") is None:
        return None

    with tempfile.TemporaryDirectory() as td:
        mini_docx = os.path.join(td, "mini.docx")
        write_single_paragraph_docx(src_paragraph, mini_docx)

        try:
            res = subprocess.run(
                ["pandoc", mini_docx, "-f", "docx", "-t", "latex", "--wrap=none"],
                capture_output=True,
                text=True,
                check=False
            )
        except Exception:
            return None

        if res.returncode != 0:
            return None

        frag = (res.stdout or "").strip()
        return extract_math_from_pandoc_latex(frag)

def latex_math_to_png(math_latex: str, out_png_path: str) -> bool:
    """
    Renders display math to PNG using LaTeX -> PDF -> PNG.
    Uses xelatex/pdflatex and either magick or pdftocairo.
    """
    latex_engine = "xelatex" if which("xelatex") else ("pdflatex" if which("pdflatex") else None)
    if latex_engine is None:
        return False

    magick = which("magick")
    pdftocairo = which("pdftocairo")

    # ensure output dir
    os.makedirs(os.path.dirname(out_png_path), exist_ok=True)

    with tempfile.TemporaryDirectory() as td:
        tex_path = os.path.join(td, "eq.tex")
        pdf_path = os.path.join(td, "eq.pdf")

        tex = r"""
\documentclass[border=2pt]{standalone}
\usepackage{amsmath,amssymb}
\begin{document}
\[
%s
\]
\end{document}
""".strip() % (math_latex or "")

        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(tex)

        # compile to pdf
        res = subprocess.run(
            [latex_engine, "-interaction=nonstopmode", "-halt-on-error", tex_path],
            cwd=td,
            capture_output=True,
            text=True
        )
        if res.returncode != 0 or (not os.path.exists(pdf_path)):
            return False

        # pdf -> png
        if magick:
            # ImageMagick: magick -density 300 eq.pdf -trim +repage out.png
            res2 = subprocess.run(
                ["magick", "-density", "300", pdf_path, "-trim", "+repage", out_png_path],
                capture_output=True,
                text=True
            )
            return res2.returncode == 0 and os.path.exists(out_png_path)

        if pdftocairo:
            # pdftocairo -png -r 300 eq.pdf outprefix  => outprefix-1.png
            outprefix = os.path.splitext(out_png_path)[0]
            res2 = subprocess.run(
                ["pdftocairo", "-png", "-r", "300", pdf_path, outprefix],
                capture_output=True,
                text=True
            )
            produced = outprefix + "-1.png"
            if res2.returncode == 0 and os.path.exists(produced):
                # rename to target name
                try:
                    os.replace(produced, out_png_path)
                except Exception:
                    shutil.copyfile(produced, out_png_path)
                return os.path.exists(out_png_path)

        return False


# ================== Language ==================
def _t(lang: str, tr: str, en: str) -> str:
    return tr if lang == "tr" else en

def ask_language() -> str:
    while True:
        lang = input("Select language / Dil seç (tr/en): ").strip().lower()
        if lang in ("tr", "en"):
            return lang
        print("Please type 'tr' or 'en'.")

def ensure_ext(path_str: str, ext: str) -> str:
    """
    Kullanıcı uzantı yazmadıysa otomatik ekler.
    ext: ".docx" gibi nokta dahil gelmeli.
    """
    s = (path_str or "").strip().strip('"').strip("'")
    if not s:
        return s
    root, current_ext = os.path.splitext(s)
    if current_ext == "":
        return s + ext
    return s

def ask_docx_path(lang: str) -> str:
    p = input(_t(
        lang,
        "Düzenlenecek .docx dosyasının adı: (.docx yazmana gerek yok) ",
        "Name of the .docx file to process: (no need to type .docx) "
    )).strip()
    p = ensure_ext(p, ".docx")
    return os.path.abspath(p)

# ================== Feature Flags ==================
@dataclass
class Features:
    use_figures: bool = True
    use_tables: bool = True
    use_equations: bool = True
    use_bibliography: bool = True
    use_plots: bool = True   # sadece assets/plots'tan ekleme (üretim yok)

def ask_feature(prompt: str, lang: str) -> bool:
    """
    TR:
      - Boş = EVET.
      - 'no' = HAYIR.
    EN:
      - Blank = YES.
      - Type 'no' = NO.
    """
    suffix = _t(lang, " (✅ = boş, ❌ = no ): ", " (✅ = blank, ❌ = no ): ")
    ans = input(prompt + suffix).strip().lower()
    return False if ans == "no" else True

def ask_features(lang: str) -> Features:
    print("\n" + _t(lang, "=== Özellik Seçimi ===", "=== Feature Selection ==="))
    use_figures = ask_feature(_t(lang, "Görsel kullanacak mısın?", "Will you use figures?"), lang)
    use_tables = ask_feature(_t(lang, "Tablo kullanacak mısın?", "Will you use tables?"), lang)
    use_equations = ask_feature(_t(lang, "Denklem ($$ ... $$) kullanacak mısın?", "Will you use equations ($$ ... $$)?"), lang)
    use_bibliography = ask_feature(_t(lang, "Kaynakça modunu kullanacak mısın?", "Will you use bibliography mode?"), lang)
    use_plots = ask_feature(_t(lang, "Grafik eklemek istiyor musun? ($grafik$)", "Do you want to insert plots? ($plot$)"), lang)

    print(_t(lang, "=======================================\n", "=======================================\n"))
    return Features(
        use_figures=use_figures,
        use_tables=use_tables,
        use_equations=use_equations,
        use_bibliography=use_bibliography,
        use_plots=use_plots,
    )

# ================== Invisible Unicode Cleaner ==================
def strip_invisible(s: str) -> str:
    """
    Word/PDF kopyala-yapıştır ile gelen görünmez karakterleri temizler.
    - NBSP (U+00A0) -> normal boşluk
    - Unicode category Cf (format chars): U+200B (zero-width space) vb. -> sil
    """
    if not s:
        return ""
    s = s.replace("\u00a0", " ")
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Cf")
    return s

def split_manual_linebreak_paragraph(raw_text: str) -> list[str]:
    """
    Word'de Shift+Enter ile aynı paragraf içinde oluşturulan satır sonlarını (\n)
    ayrı paragraflara böl.
    """
    if not raw_text:
        return []
    parts = [strip_invisible(p).strip() for p in raw_text.splitlines()]
    return [p for p in parts if p]

# ================== LaTeX Helpers ==================
def escape_latex(s: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
        "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    for k, v in replacements.items():
        s = s.replace(k, v)
    return s

def tr_upper(s: str) -> str:
    mapping = str.maketrans({
        "i": "İ", "ı": "I",
        "ş": "Ş", "ğ": "Ğ", "ü": "Ü", "ö": "Ö", "ç": "Ç",
    })
    return s.translate(mapping).upper()

def en_upper(s: str) -> str:
    return (s or "").upper()

def make_heading_detector(lang: str):
    upper_fn = tr_upper if lang == "tr" else en_upper

    def is_all_caps_heading_line(text: str) -> bool:
        t = strip_invisible(text or "").strip()
        if not t or t == "---":
            return False
        if not re.search(r"[A-Za-zÇĞİÖŞÜçğıöşü]", t):
            return False
        return upper_fn(t) == t

    def normalize_heading_text(text: str) -> str:
        return upper_fn(strip_invisible(text or "").strip())

    return is_all_caps_heading_line, normalize_heading_text

def ensure_prev_sentence_ends_with_period(latex_output, last_kind: str):
    if last_kind != "text":
        return latex_output

    for idx in range(len(latex_output) - 1, -1, -1):
        line = latex_output[idx].rstrip()
        if not line:
            continue
        if line.lstrip().startswith("\\"):
            return latex_output
        if line.endswith((".", "!", "?", ":", ";")):
            return latex_output
        latex_output[idx] = line + "."
        return latex_output

    return latex_output

# ================== CAPTION FORMATTERS ==================
def _capitalize_first(s: str) -> str:
    s = strip_invisible(s or "").strip()
    if not s:
        return ""
    return s[0].upper() + s[1:]

def format_figure_caption_with_section(raw_input: str, lang: str, section_no: int, fig_no: int) -> str:
    text = _capitalize_first(raw_input)
    label = "Şekil" if lang == "tr" else "Figure"
    out = f"{label} {section_no}.{fig_no}: {text}"
    if not out.endswith("."):
        out += "."
    return out

def format_plot_caption_with_section(raw_input: str, lang: str, section_no: int, plot_no: int) -> str:
    # İSTEK: Grafik başlıkları “Grafik 1.1: ...” formatında olsun
    text = _capitalize_first(raw_input)
    label = "Grafik" if lang == "tr" else "Plot"
    out = f"{label} {section_no}.{plot_no}: {text}"
    if not out.endswith("."):
        out += "."
    return out

def format_table_caption_with_section(raw: str, lang: str, section_no: int, tbl_no: int) -> str:
    text = _capitalize_first(raw)
    label = "Tablo" if lang == "tr" else "Table"
    out = f"{label} {section_no}.{tbl_no}: {text}"
    if not out.endswith("."):
        out += "."
    return out

# ================== TOC Writer (WITH PAGE NUMBERS) ==================
def write_toc_tex_with_pagenum(toc_entries, lang: str, out_path="toc.tex"):
    title = "İçindekiler" if lang == "tr" else "Index"

    lines = []
    lines.append(rf"\section*{{{title}}}")

    lines.append(r"\begingroup")
    lines.append(r"\setlength{\parindent}{0pt}")
    lines.append(r"\setlength{\parskip}{0pt}")
    lines.append(r"\noindent")

    for num, title_text in toc_entries:
        left_text = escape_latex(f"{num}. {title_text}")
        lines.append(
            rf"\hyperref[sec:{num}]{{{left_text}}}\dotfill "
            rf"\hyperref[sec:{num}]{{\pageref{{sec:{num}}}}}\\"
        )

    lines.append(r"\endgroup")
    lines.append("")

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

# ======================================================================
# $$ ... $$ equation parsing
# ======================================================================
def parse_dollars_equation_line(text: str):
    t = strip_invisible(text or "").strip()
    if not t.startswith("$$"):
        return None

    rest = t[2:]
    end = rest.find("$$")
    if end == -1:
        return None

    eq = rest[:end].strip()
    tail = rest[end + 2:].strip()

    if not eq:
        return None
    return eq, tail

def normalize_equation_for_latex(eq: str) -> str:
    if eq is None:
        return ""

    s = strip_invisible(eq)
    s = s.replace("\r", " ").replace("\n", " ").strip()

    # Word/Unicode -> LaTeX
    s = s.replace(r"\𝑝𝑖", r"\pi")
    s = s.replace("𝑝𝑖", r"\pi")
    s = s.replace("π", r"\pi")
    s = s.replace("σ", r"\sigma")
    s = s.replace("Σ", r"\Sigma")
    s = s.replace("∑", r"\sum")
    s = s.replace("μ", r"\mu")
    s = s.replace("Δ", r"\Delta")
    s = s.replace("Ω", r"\Omega")
    s = s.replace("×", r"\times")
    s = s.replace("·", r"\cdot")
    s = s.replace("−", "-")
    s = s.replace("≤", r"\leq")
    s = s.replace("≥", r"\geq")
    s = s.replace("≠", r"\neq")
    s = s.replace("≈", r"\approx")
    s = s.replace("ˆ", "^")

    s = re.sub(r"\s+", " ", s).strip()
    s = re.sub(r"(\\cdot)(?=[A-Za-z])", r"\\cdot ", s)
    s = re.sub(r"(\\times)(?=[A-Za-z])", r"\\times ", s)
    s = re.sub(r"(\\(?:pi|sigma|Sigma|sum|mu|Delta|Omega))(?=[A-Za-z])", r"\1{}", s)

    cleaned = []
    i = 0
    while i < len(s):
        ch = s[i]
        if ch == "\\":
            if i + 1 < len(s) and not (("a" <= s[i+1] <= "z") or ("A" <= s[i+1] <= "Z")):
                i += 1
                continue
            cleaned.append("\\")
            i += 1
            while i < len(s) and (("a" <= s[i] <= "z") or ("A" <= s[i] <= "Z")):
                cleaned.append(s[i])
                i += 1
            continue

        o = ord(ch)
        if 32 <= o <= 126 or ch in "{}^_ ":
            cleaned.append(ch)
        i += 1

    return "".join(cleaned).strip()

def replace_greek_unicode_after_escape(escaped_text: str) -> str:
    if not escaped_text:
        return escaped_text or ""

    greek_map = {
        "δ": r"\delta", "Δ": r"\Delta",
        "ε": r"\epsilon", "ϵ": r"\varepsilon",
        "ω": r"\omega", "Ω": r"\Omega",
        "μ": r"\mu",
        "σ": r"\sigma", "Σ": r"\Sigma",
        "θ": r"\theta", "Θ": r"\Theta",
        "λ": r"\lambda", "Λ": r"\Lambda",
        "φ": r"\phi", "Φ": r"\Phi", "ϕ": r"\varphi",
        "ρ": r"\rho",
        "ν": r"\nu",
        "α": r"\alpha", "β": r"\beta", "γ": r"\gamma",
        "κ": r"\kappa",
        "τ": r"\tau",
        "η": r"\eta",
        "ζ": r"\zeta",
        "ξ": r"\xi", "Ξ": r"\Xi",
        "ψ": r"\psi", "Ψ": r"\Psi",
        "χ": r"\chi",
        "ι": r"\iota",
        "π": r"\pi", "Π": r"\Pi",
    }

    out = escaped_text
    for uni, cmd in greek_map.items():
        out = out.replace(uni, rf"\({cmd}\)")
    return out

# ======================================================================
# Figure marker + image file resolution
# ======================================================================
def parse_figure_marker_line(text: str, lang: str) -> bool:
    """
    Görsel marker satırını yakalar.
    Word formatı: $fig$ / $şekil$ / $görsel$ / $photo$ / $image$ gibi.
    Büyük-küçük harf duyarsız.
    """
    t = strip_invisible(text or "").strip()
    if not (t.startswith("$") and t.endswith("$")):
        return False

    inner = t[1:-1].strip()
    if not inner:
        return False

    key = inner.casefold()

    if lang == "tr":
        tokens = {
            "fig", "şekil", "sekil",
            "görsel", "gorsel",
            "resim",
            "foto", "fotograf", "fotoğraf",
        }
    else:
        tokens = {
            "fig", "figure",
            "image", "picture",
            "photo",
        }

    return key in tokens



def is_word_list_paragraph(paragraph) -> tuple[bool, str]:
    """
    Returns (is_list, list_env) where list_env is 'itemize' or 'enumerate'.
    Uses Word numbering properties and (best-effort) style name to guess numbered lists.
    """
    try:
        ppr = paragraph._p.pPr
        numpr = ppr.numPr if ppr is not None else None
        is_list = numpr is not None
    except Exception:
        is_list = False

    if not is_list:
        return False, "itemize"

    # Best-effort: infer enumerate vs itemize from style name
    style_name = ""
    try:
        style_name = (paragraph.style.name or "")
    except Exception:
        style_name = ""
    sn = style_name.casefold()
    if "number" in sn or "numara" in sn or "enumer" in sn:
        return True, "enumerate"
    return True, "itemize"

def has_inline_image(paragraph) -> bool:
    """
    Paragraph içinde inline görsel var mı?
    - python-docx: InlineShape erişimi her zaman kolay değil; XML'den kontrol ediyoruz.
    """
    try:
        xml = paragraph._element.xml
    except Exception:
        return False
    return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("<v:imagedata" in xml)

def parse_near_caption_line(text: str, lang: str):
    """
    'Şekil 1: Başlık', 'Figure 2 - Title', 'Tablo 3.1 Başlık' gibi satırlardan:
      (kind, caption_text) döndürür.
    kind: "figure" | "table" | None
    caption_text: kullanıcı numarasız (biz yeniden numaralandıracağız)
    """
    t = strip_invisible(text or "").strip()
    if not t:
        return None, None

    # normalize spaces
    t2 = re.sub(r"\s+", " ", t)

    # Figure patterns
    if lang == "tr":
        fig_keys = r"(Şekil|Sekil)"
        tbl_keys = r"(Tablo)"
    else:
        fig_keys = r"(Figure|Fig)"
        tbl_keys = r"(Table|Tab)"

    # Şekil 1 / Şekil 1.2 / Şekil 1-2 vb. (numarayı umursamıyoruz, sadece ayıklıyoruz)
    fig_m = re.match(rf"^{fig_keys}\s*\d+(?:[\.\-]\d+)*\s*[:\-\.\)]?\s*(.*)$", t2, flags=re.I)
    if fig_m:
        cap = (fig_m.group(2) or "").strip() if fig_m.lastindex and fig_m.lastindex >= 2 else ""
        # group(2) bazı regex engine'lerinde farklı olabilir; güvenli alım:
        cap = fig_m.groups()[-1].strip() if fig_m.groups() else ""
        return "figure", cap

    tbl_m = re.match(rf"^{tbl_keys}\s*\d+(?:[\.\-]\d+)*\s*[:\-\.\)]?\s*(.*)$", t2, flags=re.I)
    if tbl_m:
        cap = tbl_m.groups()[-1].strip() if tbl_m.groups() else ""
        return "table", cap

    return None, None

def collect_caption_around(blocks, i: int, lang: str, want_kind: str, skip_elems: set):
    """
    blocks: list[(kind, obj)]
    i: current block index (image paragraph or table block)
    want_kind: "figure" or "table"
    skip_elems: caption paragraf elementlerini burada işaretleriz (output'a yazılmasın)

    Kural:
      - Önce bir üst paragrafı kontrol et.
      - Bulamazsa bir alt paragrafı kontrol et.
      - Bulduysa ve caption_text boşsa, bir sonraki paragrafı başlık devamı olarak al (tek paragraf).
    """
    def get_p_text(j):
        if 0 <= j < len(blocks) and blocks[j][0] == "p":
            return strip_invisible(blocks[j][1].text or "").strip()
        return ""

    # prefer previous paragraph caption, then next
    candidates = [i-1, i+1]
    for j in candidates:
        t = get_p_text(j)
        if not t:
            continue
        k, cap = parse_near_caption_line(t, lang=lang)
        if k == want_kind:
            # if caption has no title part, look at continuation (next paragraph)
            cap_text = cap
            if not cap_text:
                cont = get_p_text(j+1)
                # continuation satırı başka bir caption line değilse al
                ck, _ = parse_near_caption_line(cont, lang=lang)
                if cont and (ck is None):
                    cap_text = cont
                    skip_elems.add(blocks[j+1][1]._element)
            # mark caption line itself to skip
            skip_elems.add(blocks[j][1]._element)
            return cap_text.strip()

    return None

def resolve_image_path(img_idx: int) -> tuple[str | None, str | None]:
    exts = ["png", "jpg", "jpeg"]
    candidates = []
    for ext in exts:
        candidates.append(os.path.join("assets", f"image{img_idx}.{ext}"))
    for ext in exts:
        candidates.append(f"image{img_idx}.{ext}")

    for rel in candidates:
        if os.path.exists(rel):
            return rel.replace("\\", "/"), os.path.abspath(rel)
    return None, None

# ======================================================================
# Plot marker + plot file resolution (ONLY INSERTION)
# ======================================================================
def parse_plot_marker_line(text: str) -> bool:
    t = strip_invisible(text or "").strip()
    if not (t.startswith("$") and t.endswith("$")):
        return False
    inner = t[1:-1].strip().casefold()
    return inner in ("plot", "grafik")

def resolve_plot_path(plot_idx: int) -> tuple[str | None, str | None]:
    """
    assets/plots/plot1.png, plot2.png ... dosyalarını bulur.
    (grafik_uret.py buraya yazıyor.)
    """
    exts = ["png", "jpg", "jpeg"]
    candidates = []
    for ext in exts:
        candidates.append(os.path.join("assets", "plots", f"plot{plot_idx}.{ext}"))
    for ext in exts:
        candidates.append(f"plot{plot_idx}.{ext}")

    for rel in candidates:
        if os.path.exists(rel):
            return rel.replace("\\", "/"), os.path.abspath(rel)
    return None, None

# ======================================================================
# Caption marker + doc body iterator
# ======================================================================
def parse_caption_marker(text: str) -> str | None:
    t = strip_invisible(text or "").strip()
    if t == "---":
        return None
    m = re.match(r"^---\s*(.+?)\s*---$", t)
    if not m:
        return None
    inner = strip_invisible(m.group(1) or "").strip()
    return inner if inner else None

def is_tight_list_trigger(text: str) -> bool:
    """
    'Örneğin:' / 'Şunlardır:' gibi iki nokta ile biten ve kısa olan satırlar
    tight-list modunu başlatır.
    """
    t = strip_invisible(text or "").strip()
    if not t:
        return False
    # Çok uzun cümleleri trigger sayma (yanlış yakalamayı azaltır)
    if len(t) > 80:
        return False
    return t.endswith(":")

def is_tight_list_item(text: str) -> bool:
    """
    Madde gibi yazılmış satırları yakalar:
    - "konumda: f = 13.65 Hz,"
    - "a) ...", "1) ..." gibi
    - "- ..." gibi
    """
    t = strip_invisible(text or "").strip()
    if not t:
        return False

    # Çok uzun satırları madde sayma
    if len(t) > 120:
        return False

    patterns = [
        r"^[-–•]\s+\S+",            # - madde
        r"^\d+[\.\)]\s+\S+",        # 1) 1. ...
        r"^[a-zA-ZÇĞİÖŞÜçğıöşü][\)\.]\s+\S+",  # a) a. ...
        r"^[^:]{1,25}\s*:\s+\S+",   # kelime/etiket : değer (konumda: ...)
    ]
    return any(re.match(p, t) for p in patterns)

def iter_block_items(doc: Document):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    parent = doc.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield ("p", Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield ("tbl", Table(child, doc))

def table_to_latex_lines(table) -> list[str]:
    rows = table.rows
    if not rows:
        return []

    ncols = len(rows[0].cells) if rows[0].cells else 0
    if ncols <= 0:
        return []

    colspec = "|" + "|".join([r">{\centering\arraybackslash}X"] * ncols) + "|"

    lines = []
    lines.append(r"\centering")
    lines.append(r"\fontsize{12}{10}\selectfont")
    lines.append(r"\setlength{\tabcolsep}{5pt}")
    lines.append(r"\begin{tabularx}{\textwidth}{" + colspec + r"}")
    lines.append(r"\hline")

    for r_idx, row in enumerate(rows):
        cells = row.cells
        cell_texts = []
        for c in cells[:ncols]:
            raw = strip_invisible(c.text or "").replace("\n", " ").strip()
            esc = escape_latex(raw)
            esc = replace_greek_unicode_after_escape(esc)
            if r_idx == 0:
                esc = r"\textbf{" + esc + "}"
            cell_texts.append(esc)

        lines.append(" " + " & ".join(cell_texts) + r" \\")
        lines.append(r"\hline")

    lines.append(r"\end{tabularx}")
    return lines

# ================== APPENDIX / KAYNAKÇA DETECTOR ==================
def strip_leading_numbering(s: str) -> str:
    t = strip_invisible(s or "").strip()
    t = re.sub(r"^\s*\d+\s*([.)\-:]|\s)\s*", "", t)
    return t.strip()

def is_bibliography_heading(text: str, lang: str, norm_heading_fn) -> bool:
    target = "KAYNAKÇA" if lang == "tr" else "APPENDIX"
    target_norm = norm_heading_fn(target)
    cleaned = strip_leading_numbering(text)
    cleaned_norm = norm_heading_fn(cleaned)
    return cleaned_norm == target_norm

# ================== Main Conversion ==================
def convert_docx_to_latex(docx_filename: str, lang: str, features: Features):
    docx_filename = os.path.abspath(docx_filename)
    base_dir = os.path.dirname(docx_filename) or "."
    os.chdir(base_dir)

    docx_name = os.path.basename(docx_filename)
    if not os.path.exists(docx_name):
        print(_t(lang, f"❌ {docx_name} bulunamadı! (Klasör: {base_dir})", f"❌ {docx_name} not found! (Folder: {base_dir})"))
        return

    is_heading, norm_heading = make_heading_detector(lang)

    doc = Document(docx_name)
    all_paragraphs = doc.paragraphs

    # içerik başlangıcı için --- marker
    content_start_index = 0
    for idx, p in enumerate(all_paragraphs):
        t = strip_invisible(p.text or "").strip()
        if t == "---":
            content_start_index = idx + 1
            break

    # TOC
    toc_entries = []
    counter = 0
    for p in all_paragraphs[content_start_index:]:
        t = strip_invisible(p.text or "").strip()
        if is_heading(t):
            tight_list_mode = False
            counter += 1
            toc_entries.append((counter, norm_heading(t)))

    write_toc_tex_with_pagenum(toc_entries, lang=lang, out_path="toc.tex")
    print(_t(lang, "\n === Sonuçlar ===", "\n=== Results ==="))
    print(_t(lang, f"✅ Dil = {lang}", f"✅ Lang = {lang}"))
    print(_t(lang, "✅ toc.tex yazıldı.", "✅ toc.tex written."))
    print(_t(lang, f"✅ Başlık sayısı: {len(toc_entries)}", f"✅ Title number: {len(toc_entries)}"))
    print(_t(
        lang,
        f"✅ Özellikler: görsel={features.use_figures}, tablo={features.use_tables}, denklem={features.use_equations}, kaynakça={features.use_bibliography}, grafik={features.use_plots}",
        f"✅ Features: figures={features.use_figures}, tables={features.use_tables}, equations={features.use_equations}, bib={features.use_bibliography}, plots={features.use_plots}"
    ))

    latex_output = []
    latex_output.append(r"\color{black}")

    img_counter_global = 0
    plot_counter_global = 0
    eq_counter = 0
    last_kind = "none"

    heading_counter = 0
    current_section_no = 0

    fig_in_section = 0
    plot_in_section = 0
    tbl_in_section = 0

    tight_list_mode = False

    pending_caption_for_table = None
    pending_caption_for_inline_figure = None  # 'Tablo 1: ...' / 'Şekil 2: ...' satırlarını yutmak için

    # $fig$ veya $plot$ geldi -> bir sonraki --- caption --- bunu bağlayacak
    pending_media_after_marker = None
    # {"type":"figure","img_path":"...", "section_no":1, "no":2}
    # {"type":"plot","plot_path":"...", "section_no":1, "no":3}

    in_bib_section = False
    bib_counter = 0

    content_start_para = all_paragraphs[content_start_index] if content_start_index < len(all_paragraphs) else None

    def flush_pending_media_without_caption():
        nonlocal pending_media_after_marker, latex_output, last_kind
        if not pending_media_after_marker:
            return
        m = pending_media_after_marker
        pending_media_after_marker = None

        if m["type"] == "figure":
            print(_t(lang, "⚠️ UYARI: $fig$ bulundu ama altında --- caption --- yok. Görsel captionsız basıldı.",
                       "⚠️ WARNING: Found $fig$ marker but no --- caption --- below it. Inserted without caption."))
            if last_kind == "heading":
                latex_output.append(r"\vspace{\baselineskip}")
            latex_output.append("\n\\begin{figure}[H]")
            latex_output.append("  \\centering")
            latex_output.append(f"  \\includegraphics[width=0.5\\textwidth]{{{m['img_path']}}}")
            latex_output.append("\\end{figure}\n")
            last_kind = "figure"

        elif m["type"] == "plot":
            print(_t(lang, "⚠️ UYARI: $plot$ bulundu ama altında --- caption --- yok. Grafik captionsız basıldı.",
                       "⚠️ WARNING: Found $plot$ marker but no --- caption --- below it. Inserted without caption."))
            latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)
            if last_kind == "heading":
                latex_output.append(r"\vspace{\baselineskip}")
            latex_output.append("\n\\begin{figure}[H]")
            latex_output.append("  \\centering")
            latex_output.append(f"  \\includegraphics[width=0.7\\textwidth]{{{m['plot_path']}}}")
            latex_output.append("\\end{figure}\n")
            last_kind = "figure"

    blocks = list(iter_block_items(doc))

    # içerik başlangıcına kadar olan blokları atla
    start_block_index = 0
    if content_start_para is not None:
        for bi, (k, o) in enumerate(blocks):
            if k == "p" and o._element is content_start_para._element:
                start_block_index = bi
                break

    # Görsel/Tablo caption satırlarını output'a basmamak için
    skip_elems: set = set()

    # Word bullet/numbered list support (preserve • / numbering as LaTeX itemize/enumerate)
    word_list_open = False
    word_list_env = "itemize"

    def close_word_list():
        nonlocal word_list_open, word_list_env, latex_output, last_kind
        if word_list_open:
            latex_output.append(rf"\end{{{word_list_env}}}")
            latex_output.append(r"\end{samepage}")  
            word_list_open = False
            word_list_env = "itemize"
            last_kind = "text"


    for bi in range(start_block_index, len(blocks)):
        kind, obj = blocks[bi]

        if kind == "p":
            if obj._element in skip_elems:
                continue

            raw_text = strip_invisible(obj.text or "")
            text = raw_text.strip()

            # --- NEW: OMML equation paragraphs may have empty text ---
            if text == "":
                if (not in_bib_section) and features.use_equations and has_omml(obj):
                    flush_pending_media_without_caption()

                    # Try: OMML -> LaTeX math via pandoc
                    math_latex = pandoc_docx_paragraph_to_latex_math(obj)

                    eq_counter += 1
                    latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)

                    if last_kind in ("heading", "table", "figure"):
                        latex_output.append(r"\par")

                    # If pandoc worked, write as real LaTeX equation
                    if math_latex:
                        latex_output.append(r"\begin{equation}")
                        latex_output.append(rf"\tag{{{eq_counter}}}")
                        latex_output.append(math_latex)
                        latex_output.append(r"\end{equation}")
                        latex_output.append(r"\par")
                        last_kind = "equation"
                    else:
                        # Fallback: render to image and embed
                        eq_dir = os.path.join("assets", "equations")
                        os.makedirs(eq_dir, exist_ok=True)
                        png_name = f"eq_{eq_counter:03d}.png"
                        png_path = os.path.join(eq_dir, png_name).replace("\\", "/")

                        # If pandoc failed, we can still try to get something printable:
                        # render a placeholder (or empty) if no math available
                        ok = latex_math_to_png(r"\text{[Equation]}", os.path.join(eq_dir, png_name))

                        latex_output.append(r"\begin{equation}")
                        latex_output.append(rf"\tag{{{eq_counter}}}")

                        if ok:
                            latex_output.append(r"\makebox[\linewidth]{\includegraphics[height=1.2cm]{" + png_path + r"}}")
                        else:
                            latex_output.append(r"\text{[Equation rendering failed]}")
                        latex_output.append(r"\end{equation}")
                        latex_output.append(r"\par")
                        last_kind = "equation"

                elif (not in_bib_section) and features.use_figures and has_inline_image(obj):
                    # Inline görsel paragrafı: caption'ı üst/alt satırdan yakala (Şekil/Figure ...)
                    flush_pending_media_without_caption()

                    img_counter_global += 1
                    fig_in_section += 1
                    if current_section_no == 0:
                        current_section_no = 1

                    latex_img_path, _ = resolve_image_path(img_counter_global)
                    if latex_img_path is None:
                        print(_t(lang, f"⚠️ UYARI: image{img_counter_global} bulunamadı (assets/ veya kök). Figür atlandı.",
                                   f"⚠️ WARNING: image{img_counter_global} not found (assets/ or root). Figure skipped."))
                        continue

                    cap_text = None
                    if pending_caption_for_inline_figure:
                        cap_text = pending_caption_for_inline_figure
                        pending_caption_for_inline_figure = None
                    else:
                        cap_text = collect_caption_around(blocks, bi, lang=lang, want_kind="figure", skip_elems=skip_elems)
                    if cap_text:
                        cap_line = format_figure_caption_with_section(
                            cap_text, lang=lang, section_no=current_section_no, fig_no=fig_in_section
                        )
                        if last_kind == "heading":
                            latex_output.append(r"\vspace{\baselineskip}")
                        latex_output.append("\n\\begin{figure}[H]")
                        latex_output.append("  \\centering")
                        latex_output.append(f"  \\includegraphics[width=0.5\\textwidth]{{{latex_img_path}}}")
                        latex_output.append(f"  \\caption*{{{escape_latex(cap_line)}}}")
                        latex_output.append("\\end{figure}\n")
                        last_kind = "figure"
                    else:
                        # caption yakalanamadıysa captionsız bas
                        if last_kind == "heading":
                            latex_output.append(r"\vspace{\baselineskip}")
                        latex_output.append("\n\\begin{figure}[H]")
                        latex_output.append("  \\centering")
                        latex_output.append(f"  \\includegraphics[width=0.5\\textwidth]{{{latex_img_path}}}")
                        latex_output.append("\\end{figure}\n")
                        last_kind = "figure"
                    continue
                else:
                    continue




            # Word bullet/numbered lists (•) -> LaTeX itemize/enumerate
            is_list_para, list_env = (False, "itemize")
            if (not in_bib_section) and text and text != "---":
                is_list_para, list_env = is_word_list_paragraph(obj)

            if is_list_para:
                flush_pending_media_without_caption()

                # if list type changes, close previous list first
                if (not word_list_open) or (word_list_env != list_env):
                    close_word_list()
                    latex_output.append(r"\begin{samepage}")
                    latex_output.append(rf"\begin{{{list_env}}}")
                    word_list_open = True
                    word_list_env = list_env

                item_txt = replace_greek_unicode_after_escape(escape_latex(text))
                latex_output.append(rf"\item {item_txt}")
                last_kind = "text"
                continue
            else:
                # leaving a Word list
                close_word_list()

            # CAPTION MARKER (--- ... ---)
            cap_inner = parse_caption_marker(text)
            if cap_inner:
                # Eğer bir önceki satır $fig$ / $plot$ ise => bu caption o medyaya ait
                if pending_media_after_marker is not None:
                    m = pending_media_after_marker
                    pending_media_after_marker = None

                    if m["type"] == "figure":
                        cap_line = format_figure_caption_with_section(
                            cap_inner, lang=lang, section_no=m["section_no"], fig_no=m["no"]
                        )
                        if last_kind == "heading":
                            latex_output.append(r"\vspace{\baselineskip}")
                        latex_output.append("\n\\begin{figure}[H]")
                        latex_output.append("  \\centering")
                        latex_output.append(f"  \\includegraphics[width=0.5\\textwidth]{{{m['img_path']}}}")
                        latex_output.append(f"  \\caption*{{{escape_latex(cap_line)}}}")
                        latex_output.append("\\end{figure}\n")
                        last_kind = "figure"
                        continue

                    if m["type"] == "plot":
                        cap_line = format_plot_caption_with_section(
                            cap_inner, lang=lang, section_no=m["section_no"], plot_no=m["no"]
                        )
                        latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)
                        if last_kind == "heading":
                            latex_output.append(r"\vspace{\baselineskip}")
                        latex_output.append("\n\\begin{figure}[H]")
                        latex_output.append("  \\centering")
                        latex_output.append(f"  \\includegraphics[width=0.7\\textwidth]{{{m['plot_path']}}}")
                        latex_output.append(f"  \\caption*{{{escape_latex(cap_line)}}}")
                        latex_output.append("\\end{figure}\n")
                        last_kind = "figure"
                        continue

                # tablo caption'ı (eski mantık)
                if not features.use_tables:
                    continue
                if pending_caption_for_table is not None:
                    print(_t(lang,
                             f"⚠️ UYARI: Önceki tablo caption kullanılmadan yenisi geldi. Önceki atlandı: '{pending_caption_for_table}'",
                             f"⚠️ WARNING: New table caption came before the previous one was used. Skipped: '{pending_caption_for_table}'"))
                pending_caption_for_table = cap_inner
                continue

            # PLOT MARKER ($plot$ / $grafik$) => SADECE assets/plots'tan EKLE
            if parse_plot_marker_line(text):
                if not features.use_plots or in_bib_section:
                    flush_pending_media_without_caption()
                    continue

                flush_pending_media_without_caption()

                plot_counter_global += 1
                plot_in_section += 1
                if current_section_no == 0:
                    current_section_no = 1

                latex_plot_path, _ = resolve_plot_path(plot_counter_global)
                if latex_plot_path is None:
                    print(_t(
                        lang,
                        f"⚠️ UYARI: plot{plot_counter_global}.png bulunamadı (assets/plots/). Grafik atlandı.",
                        f"⚠️ WARNING: plot{plot_counter_global}.png not found (assets/plots/). Plot skipped."
                    ))
                    continue

                pending_media_after_marker = {
                    "type": "plot",
                    "plot_path": latex_plot_path,
                    "section_no": current_section_no,
                    "no": plot_in_section,
                }
                continue

            # FIGURE MARKER ($fig$ / $şekil$)
            if parse_figure_marker_line(text, lang):
                if not features.use_figures or in_bib_section:
                    flush_pending_media_without_caption()
                    continue

                flush_pending_media_without_caption()

                img_counter_global += 1
                fig_in_section += 1
                if current_section_no == 0:
                    current_section_no = 1

                latex_img_path, _ = resolve_image_path(img_counter_global)
                if latex_img_path is None:
                    print(_t(lang, f"⚠️ UYARI: image{img_counter_global} bulunamadı (assets/ veya kök). Figür atlandı.",
                               f"⚠️ WARNING: image{img_counter_global} not found (assets/ or root). Figure skipped."))
                    continue

                pending_media_after_marker = {
                    "type": "figure",
                    "img_path": latex_img_path,
                    "section_no": current_section_no,
                    "no": fig_in_section,
                }
                continue

            # EQUATION
            parsed = parse_dollars_equation_line(text)
            if parsed is not None:
                if in_bib_section or (not features.use_equations):
                    flush_pending_media_without_caption()
                    continue

                flush_pending_media_without_caption()

                eq_raw, tail_text = parsed
                eq = normalize_equation_for_latex(eq_raw)
                if eq:
                    eq_counter += 1
                    latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)

                    if last_kind in ("heading", "table", "figure"):
                        latex_output.append(r"\par")

                    latex_output.append(r"\begin{equation}")
                    latex_output.append(rf"\tag{{{eq_counter}}}")
                    latex_output.append(eq)
                    latex_output.append(r"\end{equation}")

                    if tail_text:
                        tail_text = strip_invisible(tail_text)
                        escaped_tail = escape_latex(tail_text)
                        escaped_tail = replace_greek_unicode_after_escape(escaped_tail)
                        latex_output.append(escaped_tail)
                        latex_output.append(r"\par")
                        latex_output.append(r"\vspace{\baselineskip}")
                        last_kind = "text"
                    else:
                        latex_output.append(r"\par")
                        last_kind = "equation"
                continue

            # HEADINGS
            if is_heading(text):
                flush_pending_media_without_caption()

                if pending_caption_for_table is not None:
                    print(_t(lang,
                             f"⚠️ UYARI: Başlık geldi ama bekleyen tablo caption vardı (kullanılmadı): '{pending_caption_for_table}'",
                             f"⚠️ WARNING: A heading appeared but there was a pending table caption (not used): '{pending_caption_for_table}'"))
                    pending_caption_for_table = None

                if features.use_bibliography:
                    in_bib_section = is_bibliography_heading(text, lang=lang, norm_heading_fn=norm_heading)
                    if in_bib_section:
                        bib_counter = 0
                else:
                    in_bib_section = False

                latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)
                if latex_output:
                    latex_output.append("\n\\clearpage\n")

                heading_counter += 1
                current_section_no = heading_counter

                # her ana başlıkta sayaçlar sıfırlansın
                fig_in_section = 0
                plot_in_section = 0
                tbl_in_section = 0

                title = norm_heading(text)
                display = f"{heading_counter}. {title}"

                latex_output.append(r"\phantomsection")
                latex_output.append(f"\\label{{sec:{heading_counter}}}")
                latex_output.append(f"\\section*{{{escape_latex(display)}}}")

                last_kind = "heading"
                continue

            
            # ----------------------------------------------------------
            # Auto-caption lines like "Tablo 1: ..." / "Şekil 2 - ..." :
            # Bu satırları normal metin olarak basma; ilgili tablo/görsele caption olarak bağla.
            # ----------------------------------------------------------
            if (not in_bib_section):
                auto_kind, auto_cap = parse_near_caption_line(text, lang=lang)

                # Tablo caption satırı
                if auto_kind == "table" and features.use_tables:
                    # Yakında bir tablo var mı? (boş paragrafları atlayarak bak)
                    has_next_table = False
                    for nb in range(bi + 1, min(len(blocks), bi + 5)):
                        nk, no = blocks[nb]
                        if nk == "p":
                            nt = strip_invisible(no.text or "").strip()
                            if nt == "":
                                continue
                            # tablo gelmeden normal bir metin başladıysa aramayı kes
                            break
                        if nk == "tbl":
                            has_next_table = True
                            break

                    if has_next_table:
                        # Bu caption satırını output'a basma
                        skip_elems.add(obj._element)

                        cap_text = (auto_cap or "").strip()
                        if not cap_text:
                            # Başlık devamı bir sonraki paragraf olabilir
                            if bi + 1 < len(blocks) and blocks[bi + 1][0] == "p":
                                cont_p = blocks[bi + 1][1]
                                cont = strip_invisible(cont_p.text or "").strip()
                                ck, _ = parse_near_caption_line(cont, lang=lang)
                                if cont and (ck is None):
                                    cap_text = cont
                                    skip_elems.add(cont_p._element)

                        # pending_caption_for_table varsa üstüne yazma (marker öncelikli)
                        if cap_text and pending_caption_for_table is None:
                            pending_caption_for_table = cap_text

                        continue

                # Şekil caption satırı (inline image'e bağlanacak)
                if auto_kind == "figure" and features.use_figures:
                    # Yakında inline görsel paragrafı var mı?
                    has_next_image = False
                    for nb in range(bi + 1, min(len(blocks), bi + 5)):
                        nk, no = blocks[nb]
                        if nk == "p":
                            nt = strip_invisible(no.text or "").strip()
                            if nt == "" and has_inline_image(no):
                                has_next_image = True
                                break
                            if nt == "":
                                continue
                            # görsel gelmeden normal metin başladıysa kes
                            break

                    if has_next_image:
                        skip_elems.add(obj._element)

                        cap_text = (auto_cap or "").strip()
                        if not cap_text:
                            if bi + 1 < len(blocks) and blocks[bi + 1][0] == "p":
                                cont_p = blocks[bi + 1][1]
                                cont = strip_invisible(cont_p.text or "").strip()
                                ck, _ = parse_near_caption_line(cont, lang=lang)
                                if cont and (ck is None):
                                    cap_text = cont
                                    skip_elems.add(cont_p._element)

                        if cap_text:
                            pending_caption_for_inline_figure = cap_text
                        continue
# NORMAL TEXT
            flush_pending_media_without_caption()

            # --- NEW: Word içi Shift+Enter satırlarını böl ---
            if raw_text and ("\n" in raw_text) and (not in_bib_section):
                parts = split_manual_linebreak_paragraph(raw_text)

                for part in parts:
                    if not part or part == "---":
                        continue

                    escaped_part = replace_greek_unicode_after_escape(
                        escape_latex(part)
                    )

                    latex_output.append(escaped_part)
                    latex_output.append(r"\par")
                    latex_output.append(r"\vspace{\baselineskip}")
                    last_kind = "text"
                continue


            if text and text != "---":
                if features.use_bibliography and in_bib_section:
                    bib_counter += 1
                    escaped_item = escape_latex(text)
                    escaped_item = replace_greek_unicode_after_escape(escaped_item)
                    latex_output.append(f"[{bib_counter}] {escaped_item}\par")
                    last_kind = "text"
                    continue

                escaped = escape_latex(text)
                escaped = replace_greek_unicode_after_escape(escaped)

                # --- Tight list modu başlatan satır mı? ---
                if is_tight_list_trigger(text):
                    escaped = escape_latex(text)
                    escaped = replace_greek_unicode_after_escape(escaped)

                    latex_output.append(escaped)
                    latex_output.append(r"\par")  # kendisi paragraf kalsın
                    latex_output.append(r"\vspace{\baselineskip}")  # trigger satırdan sonra boşluk olabilir
                    last_kind = "text"

                    tight_list_mode = True
                    continue

                # --- Tight list item mı? ---
                if tight_list_mode and is_tight_list_item(text):
                    escaped = escape_latex(text)
                    escaped = replace_greek_unicode_after_escape(escaped)

                    # Madde satırları arasında BOŞLUK YOK:
                    # \par + vspace yerine satır sonu (\\) kullanıyoruz.
                    latex_output.append(r"\noindent " + escaped + r"\\")
                    last_kind = "text"
                    continue

                # Eğer tight_list_mode açık ama artık madde değilse, modu kapat
                if tight_list_mode and (not is_tight_list_item(text)):
                    tight_list_mode = False

                # --- Normal paragraf (eski davranış) ---
                escaped = escape_latex(text)
                escaped = replace_greek_unicode_after_escape(escaped)
                latex_output.append(escaped)
                latex_output.append(r"\par")
                latex_output.append(r"\vspace{\baselineskip}")
                last_kind = "text"
                continue

        elif kind == "tbl":
            close_word_list()
            if not features.use_tables:
                if pending_caption_for_table is not None:
                    print(_t(lang,
                             f"⚠️ UYARI: Tablo kapalı ama caption + tablo bulundu. Caption atlandı: '{pending_caption_for_table}'",
                             f"⚠️ WARNING: Tables are disabled but a caption+table was found. Caption skipped: '{pending_caption_for_table}'"))
                    pending_caption_for_table = None
                continue

            if in_bib_section:
                continue

            flush_pending_media_without_caption()

            if pending_caption_for_table is None:
                # --- NEW: Tablo gördüğün yerin üst/alt satırlarından "Tablo/Table ..." caption'ını yakala ---
                near_cap = collect_caption_around(blocks, bi, lang=lang, want_kind="table", skip_elems=skip_elems)
                if near_cap:
                    pending_caption_for_table = near_cap
                else:
                    print(_t(lang,
                             "⚠️ UYARI: Caption marker olmadan tablo bulundu. Tablo atlandı. (Tablodan önce: --- ... --- yaz)",
                             "⚠️ WARNING: Found a table without a caption marker. Table skipped. (Write --- ... --- before the table)"))
                    continue



            tbl_in_section += 1
            if current_section_no == 0:
                current_section_no = 1

            cap_full = format_table_caption_with_section(
                pending_caption_for_table, lang=lang, section_no=current_section_no, tbl_no=tbl_in_section
            )
            pending_caption_for_table = None

            latex_output = ensure_prev_sentence_ends_with_period(latex_output, last_kind)

            if last_kind == "heading":
                latex_output.append(r"\vspace{\baselineskip}")

            latex_output.append("\n\\begin{table}[H]")
            latex_output.append("  \\centering")
            latex_output.append(f"  \\caption*{{{escape_latex(cap_full)}}}")
            latex_output.extend(["  " + ln if ln else "" for ln in table_to_latex_lines(obj)])
            latex_output.append("\\end{table}\n")

            last_kind = "table"

    close_word_list()
    flush_pending_media_without_caption()

    if pending_caption_for_table is not None and features.use_tables:
        print(_t(lang,
                 f"⚠️ UYARI: Tablo caption bulundu ama ardından tablo gelmedi: '{pending_caption_for_table}'",
                 f"⚠️ WARNING: Table caption found but no table followed: '{pending_caption_for_table}'"))

    with open("content.tex", "w", encoding="utf-8") as f:
        f.write("\n".join(latex_output))

    print("✅ " + _t(lang, "content.tex yazıldı.", "content.tex written."))
    print("ℹ️ " + _t(lang, "Sayfa numaraları için PDF'yi en az 2 kez derlemen gerekebilir (ilkinde ?? çıkabilir).",
                      "You may need to compile the PDF at least twice for page numbers (?? may appear the first time)."))

if __name__ == "__main__":
    lang = ask_language()
    features = ask_features(lang)
    docx_path = ask_docx_path(lang)
    convert_docx_to_latex(docx_path, lang=lang, features=features)
